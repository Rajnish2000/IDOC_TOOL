import io
import pytesseract
import cv2
import pandas as pd
# import docx 
# from docx.shared import Pt
from docx import Document 
from PIL import Image
from pix2tex.cli import LatexOCR
# from curses import ascii
import subprocess
from img2table.document import Image
from img2table.ocr import TesseractOCR
from collections import OrderedDict


#add your tesseract path here 
pytesseract.pytesseract.tesseract_cmd = 'C:\\Program Files\\Tesseract-OCR\\tesseract.exe'

myconfig = r"--psm 11 --oem 3"

def extract_text(image_path):

  #Loading image in CV2 format
  img = cv2.imread(image_path)

  #Image preprocessing
  gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
  cv2.imwrite("temp/gray.png", gray)
  blur = cv2.GaussianBlur(gray, (7,7), 0)
  cv2.imwrite("temp/blur.png", blur)

  # thresh = cv2.threshold(blur, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)[1]
  thresh = cv2.threshold(blur, 210, 230, cv2.THRESH_BINARY + cv2.THRESH_OTSU)[1]
  kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (3,13))
  dilate = cv2.dilate(thresh, kernel, iterations=1)

  #Finding Contours for creating boundaries
  cnts = cv2.findContours(dilate, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
  cnts = cnts[0] if len(cnts) == 2 else cnts[1]
  cnts = sorted(cnts, key=lambda x: cv2.boundingRect(x)[0])

  for c in cnts:
    x, y, w, h = cv2.boundingRect(c)
    if h > 50 and w > 10:
      cv2.rectangle(img, (x,y), (x+w, y+h), (36,200,12), 2)

  cv2.imwrite("temp/boxes.png", img)

  # Text extraction
  text = pytesseract.image_to_string(gray,config=myconfig)

#   print("Extracted Text:\n", text)
  # doc = docx.Document() 
  # doc.add_paragraph(text)
  # doc.save('new_file.docx')
  return img, text


def image_formula_to_text(image_path):
    img = Image('captured_image.jpg')
    model = LatexOCR()
    letx = model(img)
    print(letx)
    return letx

def get_docx(text):
    document = Document()
    newtext = ''.join(c for c in text if valid_xml_char_ordinal(c))
    document.add_paragraph(newtext)
    bio = io.BytesIO()
    # print(type(text))
    # print("bio here.")
    print(bio)
    document.save(bio)
    return bio.getvalue()

def valid_xml_char_ordinal(c):
    codepoint = ord(c)
    # conditions ordered by presumed frequency
    return (
        0x20 <= codepoint <= 0xD7FF or
        codepoint in (0x9, 0xA, 0xD) or
        0xE000 <= codepoint <= 0xFFFD or
        0x10000 <= codepoint <= 0x10FFFF
        )

# for cleaning. i will use it later.
# def clean(text):
#     return str(''.join(
#             ascii.isprint(c) and c or '?' for c in text
#             )) 


def latex_to_docx(input_file, output_file):
    try:
        # Command to run Pandoc
        subprocess.run(['pandoc', input_file, '-o', output_file], check=True)
        print(f"Conversion successful: {output_file}")
    except subprocess.CalledProcessError as e:
        print(f"Error during conversion: {e}")
    except FileNotFoundError:
        print("Pandoc is not installed or not found in PATH")   


# Creating LaTeX file

def create_Latex_File(text):
   with open('sampleLatTexFile.tex', 'w') as f:
      f.write(text)
  


def img_to_excel(img_path):
    ocr = TesseractOCR(n_threads=1, lang="eng")
    img = Image(src='captured_image.jpg')
    image_tables = img.extract_tables(ocr=ocr,implicit_rows=False,
                                      implicit_columns=False,
                                      borderless_tables=False,min_confidence=70)
    print(image_tables)
    data = extracting_table_value(image_tables[0].content.values())
    # print(data)
    header = data[0]
    data_rows = data[1:]
    df = pd.DataFrame(data_rows, columns=header)
    return df

def extracting_table_value(tables):
   values = []
   print('table data: ',*tables)

   for id_row, row in enumerate(tables):
    row_data = []
    for id_col, cell in enumerate(row):
        x1 = cell.bbox.x1
        y1 = cell.bbox.y1
        x2 = cell.bbox.x2
        y2 = cell.bbox.y2
        value = cell.value
        row_data.append(value)
    values.append(row_data)
   print('values is : ',values)
   return values 

